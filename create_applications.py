# ====================================================================
# Standard packages
    
import os
import re
import numpy as np
import pandas as pd


# Document manipulation
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
from PyPDF2 import PdfMerger, PdfWriter, PdfReader
from datetime import date
import subprocess

# Spatial
import folium

#==========================================================================================
# Path
path='/Users/balcazar/Dropbox/JOB MARKET APPLICATIONS/'
job_db = pd.read_csv(path+'job_market_db_2023.csv')

# Generate additional variables
job_db['apsa_link'] = 'https://www.apsanet.org/CAREERS/eJobs/eJobs-Online/JBctl/JobSearch/'+job_db['jobid'].astype(str)
job_db['university'] = job_db['company']
for strg in [' university ', ' college ',' of ',' the ',': ',' - ',
             'university','of','college','the',':','-']:
    for i in range(len(job_db['university'])):
        try:
            compiled = re.compile(re.escape(strg), re.IGNORECASE)
            job_db.loc[i,['university']] = compiled.sub('',job_db.loc[i,['university']][0])
        except:
            None
 
# Fixing typos
job_db['university']=job_db['university'].str.replace(' ','_') 
job_db['university']=job_db['university'].str.replace(' ','__') 
job_db['university']=job_db['university'].str.replace('-','') 
for strg in ['  ',',','University_','_University','_of','of_','of_']:
             job_db['university']=job_db['university'].str.replace(strg,'') 
             
for i in range(len(job_db['university'])):
    try:
        if (job_db.loc[i,['university']][0][-1]=='_')==True :
            job_db.loc[i,['university']]=job_db.loc[i,['university']][0][:-1]
    except:
        None
#==========================================================================================
# Obtaining links in description  
job_db['description']=job_db['description'].str.replace('"','') 

job_db['job_link'] = np.nan
job_db['job_contact'] = np.nan
for i in range(len(job_db['description'])):
    try:
        job_db.loc[i,['job_link']] = re.search("(?P<url>https?://[^\s]+)", job_db['description'].iloc[i]).group("url")
        job_db.loc[i,['job_contact']] = str(re.search("[a-z0-9\.\-+_]+@[a-z0-9\.\-+_]+\.[a-z]+", job_db['description'].iloc[i])).split("'")[1]
        
    except:
        job_db.loc[i,['job_link']] = ''
        job_db.loc[i,['job_contact']] = ''
#==========================================================================================
# Joining subfield and expertise
job_db = job_db.fillna('')
job_db['subfield']=job_db['subfield1']+"; "+job_db['subfield2']+"; "+job_db['subfield3']
job_db['expertise']=job_db['expertise1']+"; "+job_db['expertise2']+"; "+job_db['expertise3']
    
#==========================================================================================
# Select the jobs interest in
selected=job_db
selected=selected[selected['searchstatus'].str.contains('posted', case=False) ==True] # Availability
selected=selected[selected['dateavailable'].str.contains('2023|2024|2025', case=False) ==True]  # Job market year

# Topic selection
selected=selected[selected['position'].str.contains('assistant|post|open', case=False) ==True]  # Job type
selected=selected[selected['title'].str.contains('chair|dean|black|criminal|latino|law|latinx|visiting assistant', case=False) ==False]  # Job type
selected=selected[selected['department'].str.contains('criminal|african', case=False) ==False]  # Job type

# Geographical selection
selected=selected[selected['region'].str.contains('euro|asia', case=False) ==False] # Excluding europe region
selected=selected[selected['state'].str.contains('alabama|florida|kentucky|missouri', case=False) ==False] # Excluding asian region

# Ranking selection
for var in ['rank_us','rank_la','rank_rg']: #'rank_wus','rank_qs',
    selected=selected[(pd.to_numeric(selected[var])<70) | (pd.to_numeric(selected[var]).isna())]
    if var=='rank_la' or var=='rank_rg':  
        selected=selected[(pd.to_numeric(selected[var])<70) | (pd.to_numeric(selected[var]).isna())]

# Select those that have not yet expired
selected['date_aux']=selected['deadline'].str.replace('Open Until Filled','03/31/2024')    
selected['date_aux']=selected['date_aux'].astype('datetime64[ns]')
selected=selected[selected['date_aux']> '2023-09-01']
selected=selected.sort_values(by=['date_aux'])

#==========================================================================================
# Cleaning position field
replace_in_position=['Academic Positions','Associate Professor','Associate','Full','Full Professor',
                     'All','Other',
                     ':',',','/','  ']

for strg in replace_in_position:
    selected['position']=selected['position'].str.replace(strg,'')
     
selected['position']=selected['position'].str.replace('FellowshipsPost-docs','Post-doctoral') 
selected['position']=selected['position'].str.replace('Post-doctoral Post-doctoral','Post-doctoral')    
selected['position']=selected['position'].str.replace('Post-doctoralPost-doctoral','Post-doctoral')    
selected['position']=selected['position'].str.replace('Post-doctoralPost-doctoralPost-doctoral','Post-doctoral')    
selected['position']=selected['position'].str.replace(' Pre-doctoral','')    



# Dropping unnecessary columns
selected=selected.reset_index()
selected=selected.drop(['subfield1','subfield2','subfield3',
                        'expertise1','expertise2','expertise3',
                        'jobid','dateavailable','searchstatus',
                        'index','Unnamed: 0','date_aux'], axis=1)
#==========================================================================================
# Obtaining elements in description
attachments = ['CV', 'cv','C.V','c.v','curriculum','Curriculum',
               'writing sample',
               'research statement',
               'teaching statement', 'statement of teaching interests',
               'diversity statement',
               'evaluations',
               'letters','letter',
               'syllabus',
               'syllabi',
               'transcript']


selected['attach'] = ""

for i in range(len(selected['attach'])):
    for att in attachments:
        try:
            if (selected['description'].iloc[i].find(att)>-1)==True :
                
                if (att=='CV' or att=='cv' or att=='curriculum' or att=='C.V' or att=='c.v')== True :
                    selected.loc[i,['attach']] = selected.loc[i,['attach']][0] + 'Curriculum Vitae; '
                    
                
                if (att=='writing sample')== True :
                    selected.loc[i,['attach']] = selected.loc[i,['attach']][0] + 'writing sample; '
                    
                    
                if (att=='research statement')== True :
                    selected.loc[i,['attach']] = selected.loc[i,['attach']][0] +'research statement; '
                    
                    
                if (att=='teaching statement' or att=='statement of teaching interests')== True :
                    selected.loc[i,['attach']] = selected.loc[i,['attach']][0] +'teaching statement; '
                    
                    
                if (att=='diversity statement')== True :
                    selected.loc[i,['attach']] = selected.loc[i,['attach']][0] +'diversity statement; '
                    
                    
                if (att=='letter' or att=='letters')== True :
                    selected.loc[i,['attach']] = selected.loc[i,['attach']][0] +'three letters of recommendation; '
                    
                
                if (att=='evaluations')== True :
                    selected.loc[i,['attach']] = selected.loc[i,['attach']][0] +'teaching evaluations; '
                    
                
                if (att=='syllabus' or att=='syllabi')== True :
                    selected.loc[i,['attach']] = selected.loc[i,['attach']][0] +'syllabus; '
                    
                    
                if (att=='transcript')== True :
                    selected.loc[i,['attach']] = selected.loc[i,['attach']][0] +'transcripts; '
        except:
            selected.loc[i,['attach']] = ''
        
# Fixing typos
for var in ['attach','subfield','expertise']: 
    selected[var]=selected[var].str.replace(' ;','')
    for i in range(len(selected['attach'])):
        try:
            if (selected.loc[i,[var]][0][-2]==';')==True :
                selected.loc[i,[var]]=selected.loc[i,[var]][0][:-2]
        except:
            None

#==========================================================================================
# Selecting positions
#selected = selected.drop(index=[0,3,4,6,7,8,9,11,12,13,15,17,
#                                18,19,21,23,24,25,28,32], axis=0) # Drop manually - unwanted
#selected = selected.drop(index=[1,2,14,20,22,26,27,31], axis=0) # Drop manually - finished
selected.to_csv(path+'selected.csv',sep=',') # For QGIS


# Visualizing: To do
    
    
#==========================================================================================
# Building topics  pd.isna(selected['subfield'])

# IR and IPE
select_IR=selected[selected['subfield'].str.contains('international', case=False) |
                   selected['subfield'].str.contains(' ipe ', case=False) | 
                   selected['subfield'].str.contains(' ir ', case=False)  | 
                   selected['subfield'].str.contains('economy', case=False) |
                   selected['description'].str.contains('international', case=False) |
                   selected['description'].str.contains(' ipe ', case=False) | 
                   selected['description'].str.contains(' ir ', case=False)  | 
                   selected['description'].str.contains('economy', case=False) ==True] 
select_IR=select_IR[select_IR['position'].str.contains('post-doctoral', case=False) ==False]
select_IR=select_IR.reset_index()

# Comparative and IPE
select_PE=selected[selected['subfield'].str.contains('comparative', case=False) |
                   selected['subfield'].str.contains(' pe ', case=False) |
                   selected['subfield'].str.contains('economy', case=False) |
                   selected['description'].str.contains('comparative', case=False) |
                   selected['description'].str.contains(' pe ', case=False) |
                   selected['description'].str.contains('economy', case=False)  ==True] 
select_PE=select_PE[select_PE['position'].str.contains('post-doctoral', case=False) ==False]
select_PE=select_PE.reset_index()

# Public policy
select_policy=selected[selected['subfield'].str.contains('public', case=False) |
                       selected['subfield'].str.contains('public', case=False)  ==True ] 
select_policy=select_policy[select_policy['position'].str.contains('post-doctoral', case=False) ==False]
select_policy=select_policy.reset_index()

# Business
select_business=selected[selected['department'].str.contains('business', case=False) |
                         selected['department'].str.contains('business', case=False) == True ] 
select_business=select_business[select_business['position'].str.contains('post-doctoral', case=False) ==False]
#select_business=select_business.drop(labels=[2,16], axis=0)
#select_business.reset_index()

# Climate
select_climate=selected[(selected['position'].str.contains('climate change', case=False) | 
                   selected['position'].str.contains('environment', case=False) |
                   selected['description'].str.contains('climate change', case=False) | 
                   selected['description'].str.contains('environment', case=False)) &
                   ((selected['subfield'].str.contains('international ', case=False) | 
                   selected['subfield'].str.contains(' ipe ', case=False)  | 
                   selected['subfield'].str.contains(' ir ', case=False)  | 
                   selected['subfield'].str.contains(' pe ', case=False)  | 
                   selected['subfield'].str.contains('comparative', case=False)  | 
                   selected['subfield'].str.contains('economy', case=False)) |
                   (selected['description'].str.contains('international ', case=False) | 
                   selected['description'].str.contains(' ipe ', case=False)  | 
                   selected['description'].str.contains(' ir ', case=False)  | 
                   selected['description'].str.contains(' pe ', case=False)  | 
                   selected['description'].str.contains('comparative', case=False)  | 
                   selected['description'].str.contains('economy', case=False)))==True]
select_climate=select_climate[select_climate['position'].str.contains('post-doctoral', case=False) ==False]
select_climate=select_climate.reset_index()

# Postdocs
select_pdocsx=selected[selected['position'].str.contains('Post-doctoral', case=False) &
                   ((selected['subfield'].str.contains('international ', case=False) | 
                   selected['subfield'].str.contains(' ipe ', case=False)  | 
                   selected['subfield'].str.contains(' ir ', case=False)  | 
                   selected['subfield'].str.contains(' pe ', case=False)  | 
                   selected['subfield'].str.contains('comparative', case=False)  | 
                   selected['subfield'].str.contains('economy', case=False)) |
                   (selected['description'].str.contains('international ', case=False) | 
                   selected['description'].str.contains(' ipe ', case=False)  | 
                   selected['description'].str.contains(' ir ', case=False)  | 
                   selected['description'].str.contains(' pe ', case=False)  | 
                   selected['description'].str.contains('comparative', case=False)  | 
                   selected['description'].str.contains('economy', case=False))) ==True]
select_pdocsx=select_pdocsx.reset_index()
for field in ['IR','PE','policy','business','climate','pdocsx']:
    eval(f"select_{field}").to_csv(path+'selected_'+field+'.csv',sep=',') 

#==========================================================================================
# Create applications

# Date
date = date.today().strftime("%B %d, %Y")

# Run loop over different types of applications
for field in ['IR','PE','policy','business','climate','pdocsx']:
    
    # Modifying documents conditional on field
    for index, row in eval(f"select_{field}.iterrows()"):
        
        for doc in ['Cover_letter','Research_statement','Teaching_statement']: 
            #for field in ['IR','PE','policy','business','climate','pdocsx']: 
            globals()[f'{doc}_{field}'] = Document(path+'attachments/'+field+'/'+doc+'_'+field+'.docx')
            
            style = eval(f"{doc}_{field}.styles['Normal']")
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(11)

    
        try:   
            # Obtaining fields 
            date = date
            position = row['title']
            university = row['company']
            department = row['department']
            documents = row['attach']
            app = row['university']
            
            # Make directory
            if os.path.isdir(path+'/apps 2023/'+app)==False:
                os.makedirs(path+'/apps 2023/'+app)
            
            # Field specific subfield and specialization
            if field=='IR':
                subfield='International Relations'
                specialization='International Political Economy'
                
            if field=='PE':
                subfield='Comparative Politics'
                specialization='Political Economy'
                  
            if field=='policy':
                subfield='Political Economy'
                specialization='Political Economy'
                
            if field=='business':
                subfield='Political Economy'
                specialization='International Political Economy'
                
            if field=='climate':
                subfield='Political Economy'
                specialization='Political Economy'
            
            if field=='pdocsx':
                subfield='International Political Economy'
                specialization='International Political Economy'
                    
            # Cover letter    
            try:
                globals()[f"Cover_letter_{field}_m"]=eval(f"Cover_letter_{field}")
                
                # Replace date
                eval(f"Cover_letter_{field}_m").paragraphs[0].text=eval(f"Cover_letter_{field}_m").paragraphs[0].text.replace('[DATE]',date)
                # Replace position
                eval(f"Cover_letter_{field}_m").paragraphs[4].text=eval(f"Cover_letter_{field}_m").paragraphs[4].text.replace('[POSITION]',position)
                # Replace university
                eval(f"Cover_letter_{field}_m").paragraphs[4].text=eval(f"Cover_letter_{field}_m").paragraphs[4].text.replace('[UNIVERSITY]',university)
                # Replace department
                eval(f"Cover_letter_{field}_m").paragraphs[4].text=eval(f"Cover_letter_{field}_m").paragraphs[4].text.replace('[DEPARTMENT]',department)
                # Replace subfield
                eval(f"Cover_letter_{field}_m").paragraphs[4].text=eval(f"Cover_letter_{field}_m").paragraphs[4].text.replace('[SUBFIELD]',subfield)
                # Replace specialization
                eval(f"Cover_letter_{field}_m").paragraphs[4].text=eval(f"Cover_letter_{field}_m").paragraphs[4].text.replace('[SPECIALIZATION]',specialization)
                # Replace specialization
                eval(f"Cover_letter_{field}_m").paragraphs[26].text=eval(f"Cover_letter_{field}_m").paragraphs[26].text.replace('[DOCUMENTS]',documents)
                
                eval(f"Cover_letter_{field}_m").save(path+'/apps 2023/'+app+'/'+'Cover_letter_'+field+'.docx')
            
            except:
                print('Problem with application cover letter for ' + app)
                
            # Research statement  
            try:  
                globals()[f"Research_statement_{field}_m"]=eval(f"Research_statement_{field}")
                
                # Replace date
                eval(f"Research_statement_{field}_m").paragraphs[0].text=eval(f"Research_statement_{field}_m").paragraphs[0].text.replace('[DATE]',date)
                # Replace position
                eval(f"Research_statement_{field}_m").paragraphs[4].text=eval(f"Research_statement_{field}_m").paragraphs[4].text.replace('[POSITION]',position)
                # Replace university
                eval(f"Research_statement_{field}_m").paragraphs[4].text=eval(f"Research_statement_{field}_m").paragraphs[4].text.replace('[UNIVERSITY]',university)
                # Replace department
                eval(f"Research_statement_{field}_m").paragraphs[4].text=eval(f"Research_statement_{field}_m").paragraphs[4].text.replace('[DEPARTMENT]',department)
                # Replace subfield
                eval(f"Research_statement_{field}_m").paragraphs[4].text=eval(f"Research_statement_{field}_m").paragraphs[4].text.replace('[SUBFIELD]',subfield)
                # Replace specialization
                eval(f"Research_statement_{field}_m").paragraphs[4].text=eval(f"Research_statement_{field}_m").paragraphs[4].text.replace('[SPECIALIZATION]',specialization)
                # Replace specialization
                eval(f"Research_statement_{field}_m").paragraphs[26].text=eval(f"Research_statement_{field}_m").paragraphs[26].text.replace('[DOCUMENTS]',documents)
                
                eval(f"Research_statement_{field}_m").save(path+'/apps 2023/'+app+'/'+'Research_statement_'+field+'.docx')
            
            except:
                print('Problem with application research statement for ' + app)
                
            
            # Teaching statement
            try:  
                globals()[f"Teaching_statement_{field}_m"]=eval(f"Teaching_statement_{field}")
                
                # Replace date
                eval(f"Teaching_statement_{field}_m").paragraphs[0].text=eval(f"Teaching_statement_{field}_m").paragraphs[0].text.replace('[DATE]',date)
                # Replace position
                eval(f"Teaching_statement_{field}_m").paragraphs[4].text=eval(f"Teaching_statement_{field}_m").paragraphs[4].text.replace('[POSITION]',position)
                # Replace university
                eval(f"Teaching_statement_{field}_m").paragraphs[4].text=eval(f"Teaching_statement_{field}_m").paragraphs[4].text.replace('[UNIVERSITY]',university)
                # Replace department
                eval(f"Teaching_statement_{field}_m").paragraphs[4].text=eval(f"Teaching_statement_{field}_m").paragraphs[4].text.replace('[DEPARTMENT]',department)
                # Replace subfield
                eval(f"Teaching_statement_{field}_m").paragraphs[4].text=eval(f"Teaching_statement_{field}_m").paragraphs[4].text.replace('[SUBFIELD]',subfield)
                # Replace specialization
                eval(f"Teaching_statement_{field}_m").paragraphs[4].text=eval(f"Teaching_statement_{field}_m").paragraphs[4].text.replace('[SPECIALIZATION]',specialization)
                # Replace specialization
                eval(f"Teaching_statement_{field}_m").paragraphs[26].text=eval(f"Teaching_statement_{field}_m").paragraphs[26].text.replace('[DOCUMENTS]',documents)
                
                eval(f"Teaching_statement_{field}_m").save(path+'/apps 2023/'+app+'/'+'Teaching_statement_'+field+'.docx')
            
            except:
                print('Problem with application cover teaching statement for ' + app)
        
        except:
            print('Problem with application for ' + app)

#==========================================================================================
# Conver word documents into PDF
def generate_pdf(doc_path, path):
    subprocess.call(['/Applications/LibreOffice.app/Contents/MacOS/soffice',
                 "--headless", 
                 "--convert-to", 
                 "pdf", 
                 "--outdir",
                 path,
                 doc_path])
    return doc_path

folder_list=os.listdir(path+'/apps 2023/')
for folder in folder_list:
    try:
        documents=os.listdir(path+'apps 2023/'+folder)
        for d in documents:
            if d[-5:]=='.docx':
                generate_pdf(path+'/apps 2023/'+folder+'/'+d, path+'/apps 2023/'+folder)
    except:
        None






generate_pdf(path+'/apps 2023/'+folder+'/'+d, path+'/apps 2023/'+folder)



'''  
#==========================================================================================
# Select the applications

for field in ['IR','PE','policy','business','climate']:
    for index, row in eval(f"select_{field}.iterrows()"):
        app = row['university']
        if app=='Wellesley':
            for doc in ['Cover_letter','Research_statement','Teaching_statement']:
                try:
                    convert(path+'/apps 2023/'+app+'/'+doc+'_'+field+'.docx', path+'/apps 2023/'+app+'/'+doc+'_'+field+'.pdf')
                except:
                    None

              
#==========================================================================================
# MERGERS
#==========================================================================================
aux_path=path+'attachments/standard_attachments/'
#==========================================================================================
# Mergining teaching evaluations
merger = PdfMerger()
pdfs=[
     'ip_seminar_fall_2018.pdf',
     'ped_spring_2019.pdf',
     'bpe_senior_seminar_fall_2020.pdf',
     'bpe_senior_seminar_fall_2021.pdf',
     'Collected teaching evaluations for Intro to Quantitative Political Analysis II.pdf',
     'bpe_senior_seminar_fall_2022.pdf',
     'IR_senior_spring_2022.pdf',
     'IR_senior_spring_2022b.pdf'
     ]

# Merge pdfs
for pdf in pdfs:
    merger.append(aux_path+'teaching evaluations/'+pdf)

merger.write(aux_path+'teaching_evaluations.pdf')
merger.close()

#==========================================================================================
# Mergining Pdfs if necessary
merger = PdfMerger()
# Manual
pdfs = [ 'Letter_teacher_evals.pdf',
         'teaching_evaluations.pdf',
         'Letter_additional_material.pdf',
         #'Letter_additional_material_2.pdf',
         'IPE_core_Syllabus.pdf',
         'Syllabus_Fall_2022_Stern.pdf',
         'Recitation_0.pdf',
         '3_bias_fixed_effects.pdf'
        ]

# Merge pdfs
merger.append(path+'Apps 2023/SUNYat_albany/Teaching_statement_IR.pdf')
for pdf in pdfs:
    merger.append(aux_path+pdf)

merger.write(path+'Apps 2023/SUNYat_albany/teaching_portfolio.pdf')
merger.close()



#==========================================================================================
# Mergining teaching evaluations
merger = PdfMerger()
pdfs=[
      'Published paper (JDE).pdf',
     'Unions and robots.pdf',
     'Tariff revenues matter for democratization.pdf',
     ]

# Merge pdfs
for pdf in pdfs:
    merger.append(aux_path+pdf)

merger.write(aux_path+'research.pdf')
merger.close()






#==========================================================================================
# pdf fixer AREA


#pages_to_keep = list(range(1,26)) # page numbering starts from 0
#infile = PdfReader(aux_path+'Recitation_0 copy.pdf', 'rb')
#output = PdfWriter()

#for i in pages_to_keep:
#    p = infile.pages[i] 
#    output.add_page(p)

#with open(aux_path+'Recitation_0.pdf', 'wb') as f:
#    output.write(f)
'''
                
